"""
Generate the final .docx report from the assembled_report.md
Applies: Calibri 12pt, 1.5 line spacing, 2.54cm margins,
structured Word headings, formatted tables, APA references.
"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# ── Paths ──────────────────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_MD = os.path.join(SCRIPT_DIR, "assembled_report.md")
OUTPUT_DOCX = os.path.join(SCRIPT_DIR, "Strategic_Asset_Allocation_Brief.docx")

# ── Read source ────────────────────────────────────────────────────────────
with open(INPUT_MD, "r", encoding="utf-8") as f:
    raw = f.read()

# ── Critical fixes from review report ──────────────────────────────────────

# 1. Resolve [CITE NEEDED] tags
raw = raw.replace(
    "[CITE NEEDED — empirical factor decomposition for the specific asset classes]",
    "(Ang, 2014)"
)
raw = raw.replace(
    "[CITE NEEDED — regime-dependent correlation literature, e.g., Longin & Solnik, 2001]",
    "(Longin & Solnik, 2001)"
)

# 2. Full APA 7th references
APA_REFERENCES = [
    "Amihud, Y., & Mendelson, H. (1986). Asset pricing and the bid-ask spread. "
    "Journal of Financial Economics, 17(2), 223–249. https://doi.org/10.1016/0304-405X(86)90065-6",

    "Ang, A. (2014). Asset management: A systematic approach to factor investing. "
    "Oxford University Press.",

    "Ang, A., Papanikolaou, D., & Westerfield, M. M. (2014). Portfolio choice with illiquid assets. "
    "Management Science, 60(11), 2737–2761. https://doi.org/10.1287/mnsc.2014.1986",

    "Asness, C. S., Frazzini, A., & Pedersen, L. H. (2012). Leverage aversion and risk parity. "
    "Financial Analysts Journal, 68(1), 47–59. https://doi.org/10.2469/faj.v68.n1.1",

    "Asness, C. S., Moskowitz, T. J., & Pedersen, L. H. (2013). Value and momentum everywhere. "
    "The Journal of Finance, 68(3), 929–985. https://doi.org/10.1111/jofi.12021",

    "Black, F., & Litterman, R. (1992). Global portfolio optimization. "
    "Financial Analysts Journal, 48(5), 28–43. https://doi.org/10.2469/faj.v48.n5.28",

    "Brinson, G. P., Hood, L. R., & Beebower, G. L. (1986). Determinants of portfolio performance. "
    "Financial Analysts Journal, 42(4), 39–44. https://doi.org/10.2469/faj.v42.n4.39",

    "Campbell, J. Y., & Viceira, L. M. (2002). Strategic asset allocation: Portfolio choice for "
    "long-term investors. Oxford University Press.",

    "DeMiguel, V., Garlappi, L., & Uppal, R. (2009). Optimal versus naive diversification: "
    "How inefficient is the 1/N portfolio strategy? The Review of Financial Studies, 22(5), "
    "1915–1953. https://doi.org/10.1093/rfs/hhm075",

    "Dimson, E., Marsh, P., & Staunton, M. (2021). Credit Suisse global investment returns "
    "yearbook 2021. Credit Suisse Research Institute.",

    "Fama, E. F., & French, K. R. (2015). A five-factor asset pricing model. "
    "Journal of Financial Economics, 116(1), 1–22. https://doi.org/10.1016/j.jfineco.2014.10.010",

    "Ibbotson, R. G., & Kaplan, P. D. (2000). Does asset allocation policy explain 40, 90, or "
    "100 percent of performance? Financial Analysts Journal, 56(1), 26–33. "
    "https://doi.org/10.2469/faj.v56.n1.2327",

    "Ilmanen, A. (2011). Expected returns: An investor’s guide to harvesting market rewards. "
    "John Wiley & Sons.",

    "Kinlaw, W. B., Kritzman, M. P., & Turkington, D. (2017). A practitioner’s guide to asset "
    "allocation. John Wiley & Sons.",

    "Longin, F., & Solnik, B. (2001). Extreme correlation of international equity markets. "
    "The Journal of Finance, 56(2), 649–676. https://doi.org/10.1111/0022-1082.00340",

    "Markowitz, H. (1952). Portfolio selection. The Journal of Finance, 7(1), 77–91. "
    "https://doi.org/10.1111/j.1540-6261.1952.tb01525.x",

    "Qian, E. (2011). Risk parity and diversification. The Journal of Investing, 20(1), "
    "119–127. https://doi.org/10.3905/joi.2011.20.1.119",
]

# ── Remove internal assembly notes and assembler footer ────────────────────
raw = re.sub(
    r'## ASSEMBLY NOTES.*?(?=\n## 1\. )',
    '',
    raw,
    flags=re.DOTALL,
)
raw = raw.replace(
    "\n*End of assembled report. This document was compiled from `sections_1_5.md`, "
    "`sections_2_3_4.md`, and `tables.md` by the h_assembler agent.*",
    ""
)
raw = re.sub(r'\*\*Status:\*\*.*?\n', '', raw)

# ── Remove the placeholder references section from MD ──────────────────────
# We'll add proper references programmatically
raw = re.sub(
    r'## References\n.*',
    '',
    raw,
    flags=re.DOTALL,
)


# ── Helper functions ───────────────────────────────────────────────────────

def set_cell_font(cell, text, bold=False, size=10, italic=False, align='left'):
    """Set font properties for a table cell."""
    cell.text = ''
    p = cell.paragraphs[0]
    # Strip markdown bold markers
    clean = text.replace('**', '')
    run = p.add_run(clean)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold or ('**' in text)
    run.font.italic = italic
    if align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0


def add_table_borders(table):
    """Add borders to all cells in a table."""
    tbl = table._tbl
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>'
    )
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    tblPr.append(tblBorders)


def shade_header_row(table):
    """Apply dark blue shading to header row with white text."""
    for cell in table.rows[0].cells:
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="1F497D" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True


def guess_alignment(header):
    """Guess cell alignment based on header."""
    numeric_kw = ['return', 'volatility', 'drawdown', 'sharpe', 'ratio',
                  'spread', 'score', 'weight', 'contribution', 'bps', '%']
    h_lower = header.lower()
    for kw in numeric_kw:
        if kw in h_lower:
            return 'right'
    if header.strip() in ('#', ''):
        return 'center'
    # Correlation matrix numeric cells
    if header.strip().startswith('Class'):
        return 'center'
    return 'left'


def parse_md_table_block(lines):
    """Parse a block of markdown table lines into (headers, rows).
    Handles tables where header cells may contain bold markdown."""
    if len(lines) < 2:
        return None, None

    def split_row(line):
        # Remove leading/trailing pipes and split
        line = line.strip()
        if line.startswith('|'):
            line = line[1:]
        if line.endswith('|'):
            line = line[:-1]
        return [c.strip() for c in line.split('|')]

    # Find header line (first non-separator pipe line)
    header_idx = None
    sep_idx = None
    for idx, l in enumerate(lines):
        stripped = l.strip()
        if re.match(r'^\|[\s:|-]+\|$', stripped):
            sep_idx = idx
        elif stripped.startswith('|') and header_idx is None:
            header_idx = idx

    if header_idx is None:
        return None, None

    headers = split_row(lines[header_idx])

    # Data rows: everything after the separator (or after header if no separator)
    start = (sep_idx + 1) if sep_idx is not None else (header_idx + 1)
    rows = []
    for l in lines[start:]:
        stripped = l.strip()
        if stripped.startswith('|') and not re.match(r'^\|[\s:|-]+\|$', stripped):
            row = split_row(stripped)
            rows.append(row)

    return headers, rows


def add_formatted_table(doc, headers, rows, caption=None, notes=None):
    """Add a formatted Word table with borders and header shading."""
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, h in enumerate(headers):
        clean = h.replace('**', '')
        set_cell_font(table.rows[0].cells[i], clean, bold=True, size=10, align='center')

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx in range(min(len(row), ncols)):
            val = row[c_idx]
            align = guess_alignment(headers[c_idx])
            # Bold detection
            is_bold = '**' in val
            clean_val = val.replace('**', '')
            set_cell_font(table.rows[r_idx + 1].cells[c_idx], clean_val,
                          bold=is_bold, size=10, align=align)

    add_table_borders(table)
    shade_header_row(table)

    if notes:
        p = doc.add_paragraph()
        run = p.add_run(notes)
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.italic = True
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(8)

    return table


def add_rich_paragraph(doc, text, style_name='Normal'):
    """Add a paragraph with inline **bold** and *italic* formatting."""
    p = doc.add_paragraph(style=style_name)
    # Split on bold and italic markers
    parts = re.split(r'(\*\*.*?\*\*|\*[^*]+?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            # Handle backtick code spans
            code_parts = re.split(r'(`[^`]+`)', part)
            for cp in code_parts:
                if cp.startswith('`') and cp.endswith('`'):
                    run = p.add_run(cp[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(11)
                else:
                    run = p.add_run(cp)
        run.font.name = run.font.name or 'Calibri'
        if not run.font.size:
            run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    return p


# ── Build document ─────────────────────────────────────────────────────────
doc = Document()

# ── Page setup ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Default styles ─────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(12)
style.font.color.rgb = RGBColor(0, 0, 0)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)

for level, size, color in [
    ('Heading 1', 18, RGBColor(0, 0, 0)),
    ('Heading 2', 14, RGBColor(0x1F, 0x49, 0x7D)),
    ('Heading 3', 12, RGBColor(0x1F, 0x49, 0x7D)),
]:
    hs = doc.styles[level]
    hs.font.name = 'Calibri'
    hs.font.size = Pt(size)
    hs.font.bold = True
    hs.font.color.rgb = color
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)
    hs.paragraph_format.line_spacing = 1.5

# List Number style
if 'List Number' in doc.styles:
    ln = doc.styles['List Number']
    ln.font.name = 'Calibri'
    ln.font.size = Pt(12)
    ln.paragraph_format.line_spacing = 1.5

# ── Parse content into blocks ──────────────────────────────────────────────
lines = raw.split('\n')
blocks = []  # list of (type, content) tuples

i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # Skip empty lines
    if not stripped:
        i += 1
        continue

    # Horizontal rules → skip
    if stripped == '---':
        i += 1
        continue

    # H1
    if stripped.startswith('# ') and not stripped.startswith('## '):
        blocks.append(('h1', stripped[2:].strip()))
        i += 1
        continue

    # H2
    if stripped.startswith('## ') and not stripped.startswith('### '):
        blocks.append(('h2', stripped[3:].strip()))
        i += 1
        continue

    # H3
    if stripped.startswith('### '):
        blocks.append(('h3', stripped[4:].strip()))
        i += 1
        continue

    # Metadata lines
    if stripped.startswith('**Course:**') or stripped.startswith('**Date:**'):
        blocks.append(('meta', stripped.replace('**', '')))
        i += 1
        continue

    # Table caption
    if re.match(r'^\*\*Table \d+:', stripped):
        caption = stripped.replace('**', '')
        # Now collect the table lines
        i += 1
        # Skip blank lines between caption and table
        while i < len(lines) and not lines[i].strip():
            i += 1
        # Collect table lines
        tbl_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            tbl_lines.append(lines[i])
            i += 1
        # Collect notes (lines starting with *)
        note_parts = []
        while i < len(lines):
            nl = lines[i].strip()
            if nl.startswith('*') and not nl.startswith('**'):
                # Clean the note text
                clean_note = nl
                if clean_note.startswith('*') and not clean_note.startswith('**'):
                    clean_note = clean_note[1:]
                if clean_note.endswith('*'):
                    clean_note = clean_note[:-1]
                note_parts.append(clean_note.strip())
                i += 1
            elif not nl:
                i += 1
                # Check if next line is also a note
                if i < len(lines) and lines[i].strip().startswith('*') and not lines[i].strip().startswith('**'):
                    continue
                break
            else:
                break
        notes = ' '.join(note_parts) if note_parts else None
        blocks.append(('table', (caption, tbl_lines, notes)))
        continue

    # Standalone table (no caption — like correlation matrix if caption was separate)
    if stripped.startswith('|') and '|' in stripped[1:]:
        tbl_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            tbl_lines.append(lines[i])
            i += 1
        # Collect notes
        note_parts = []
        while i < len(lines):
            nl = lines[i].strip()
            if nl.startswith('*') and not nl.startswith('**'):
                clean_note = nl.lstrip('*').rstrip('*').strip()
                note_parts.append(clean_note)
                i += 1
            elif not nl:
                i += 1
                if i < len(lines) and lines[i].strip().startswith('*') and not lines[i].strip().startswith('**'):
                    continue
                break
            else:
                break
        notes = ' '.join(note_parts) if note_parts else None
        blocks.append(('table', (None, tbl_lines, notes)))
        continue

    # Numbered list
    if re.match(r'^\d+\.\s', stripped):
        items = [stripped]
        i += 1
        while i < len(lines):
            nl = lines[i].strip()
            if re.match(r'^\d+\.\s', nl):
                items.append(nl)
                i += 1
            elif not nl:
                i += 1
                if i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                    continue
                break
            else:
                break
        blocks.append(('list', items))
        continue

    # Bullet list
    if stripped.startswith('- '):
        items = [stripped[2:]]
        i += 1
        while i < len(lines):
            nl = lines[i].strip()
            if nl.startswith('- '):
                items.append(nl[2:])
                i += 1
            elif not nl:
                i += 1
                if i < len(lines) and lines[i].strip().startswith('- '):
                    continue
                break
            else:
                break
        blocks.append(('bullets', items))
        continue

    # Regular paragraph
    if stripped:
        blocks.append(('para', stripped))

    i += 1

# ── Render blocks to document ──────────────────────────────────────────────
for btype, content in blocks:
    if btype == 'h1':
        h = doc.add_heading(content, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    elif btype == 'meta':
        p = doc.add_paragraph()
        run = p.add_run(content)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)

    elif btype == 'h2':
        doc.add_heading(content, level=2)

    elif btype == 'h3':
        doc.add_heading(content, level=3)

    elif btype == 'table':
        caption, tbl_lines, notes = content
        headers, rows = parse_md_table_block(tbl_lines)
        if headers and rows:
            add_formatted_table(doc, headers, rows, caption=caption, notes=notes)
        elif headers:
            # Table with headers but no data rows — still render
            add_formatted_table(doc, headers, [], caption=caption, notes=notes)

    elif btype == 'list':
        for item in content:
            text = re.sub(r'^\d+\.\s*', '', item)
            p = add_rich_paragraph(doc, text, 'List Number')

    elif btype == 'bullets':
        for item in content:
            p = add_rich_paragraph(doc, item, 'List Bullet')

    elif btype == 'para':
        add_rich_paragraph(doc, content)

# ── References ─────────────────────────────────────────────────────────────
doc.add_page_break()
doc.add_heading('References', level=2)

for ref in APA_REFERENCES:
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(ref)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)

# ── Page numbers ───────────────────────────────────────────────────────────
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)
    for r in [run, run2, run3]:
        r.font.name = 'Calibri'
        r.font.size = Pt(10)

# ── Header with course info ───────────────────────────────────────────────
for section in doc.sections:
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('QBUS3600 — Strategic Asset Allocation Brief')
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.italic = True

# ── Save ───────────────────────────────────────────────────────────────────
doc.save(OUTPUT_DOCX)
print(f"✅ Document saved to: {OUTPUT_DOCX}")
print(f"   File size: {os.path.getsize(OUTPUT_DOCX):,} bytes")

# Verify structure
doc2 = Document(OUTPUT_DOCX)
headings = [(p.style.name, p.text[:60]) for p in doc2.paragraphs if p.style.name.startswith('Heading')]
print(f"\n=== Headings ({len(headings)}) ===")
for sname, txt in headings:
    print(f"  [{sname}] {txt}")
print(f"\n=== Tables: {len(doc2.tables)} ===")
for idx, t in enumerate(doc2.tables):
    hdr = [c.text for c in t.rows[0].cells]
    print(f"  Table {idx+1}: {len(t.rows)}r x {len(t.columns)}c | Header: {hdr[:4]}...")
print(f"\n=== References: {len(APA_REFERENCES)} APA entries ===")
print(f"=== Format: Calibri 12pt, 1.5 spacing, 2.54cm margins ===")
