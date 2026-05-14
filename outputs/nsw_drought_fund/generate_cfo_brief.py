#!/usr/bin/env python3
"""Generate the CFO Briefing .docx — 2 pages, formatted per spec."""

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy


def set_cell_shading(cell, color_hex):
    """Set cell background shading."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_spacing(paragraph, space_before=0, space_after=0, line_spacing=1.15):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing


def add_bold_run(paragraph, text, font_size=12, font_name="Calibri"):
    """Add a bold run to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    run.font.name = font_name
    return run


def add_run(paragraph, text, bold=False, font_size=12, font_name="Calibri"):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = font_name
    return run


def format_table_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                       font_size=10, font_name="Calibri"):
    """Format a table cell with text."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = font_name
    # Reduce cell paragraph spacing
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.0


def set_table_borders(table):
    """Set thin borders on entire table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def create_cfo_brief():
    doc = Document()

    # ── Global style defaults ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    # Heading styles
    for level, size in [(1, 16), (2, 14), (3, 12)]:
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Calibri"
        hs.font.size = Pt(size)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)  # Navy
        hs.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        hs.paragraph_format.space_after = Pt(6)
        hs.paragraph_format.line_spacing = 1.15

    # List Bullet style
    lb = doc.styles["List Bullet"]
    lb.font.name = "Calibri"
    lb.font.size = Pt(12)
    lb.paragraph_format.line_spacing = 1.15
    lb.paragraph_format.space_after = Pt(4)

    # ── Page setup: portrait, 2.54 cm margins ──
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    margin = Cm(2.54)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin

    # ── Header ──
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("NSW Drought Fund  |  Confidential")
    hr.font.size = Pt(8)
    hr.font.name = "Calibri"
    hr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ── Footer with page numbers ──
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Page ")
    fr.font.size = Pt(8)
    fr.font.name = "Calibri"
    fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    # PAGE field
    fld_char1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fld_char2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run2 = fp.add_run()
    run2._r.append(fld_char1)
    run3 = fp.add_run()
    run3._r.append(instr)
    run4 = fp.add_run()
    run4._r.append(fld_char2)
    fr2 = fp.add_run(" of 2")
    fr2.font.size = Pt(8)
    fr2.font.name = "Calibri"
    fr2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ═══════════════════════════════════════════════════
    # PAGE 1
    # ═══════════════════════════════════════════════════

    # Title
    title = doc.add_heading("CFO Briefing — NSW Drought Fund: Initial Portfolio Recommendation", level=1)
    title.runs[0].font.size = Pt(16)
    set_paragraph_spacing(title, space_before=0, space_after=12)

    # Subtitle line
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sr = sub.add_run("Prepared for the Board of the NSW Drought Fund  |  May 2026")
    sr.font.size = Pt(10)
    sr.font.name = "Calibri"
    sr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    set_paragraph_spacing(sub, space_before=0, space_after=12)

    # ── Section 1: What ──
    doc.add_heading("Section 1 — What", level=2)

    p1 = doc.add_paragraph(style="List Bullet")
    add_run(p1, "We recommend allocating ")
    add_bold_run(p1, "15% to STI, 35% to MTG, and 50% to LTG")
    add_run(p1, ". Under a disinflationary base case — Australian CPI at 2.7% with gradual RBA easing — the mix tilts toward long-term growth, delivering a forecast gross return of ")
    add_bold_run(p1, "8.7% p.a.")
    add_run(p1, " at estimated portfolio risk of ")
    add_bold_run(p1, "~8.1%")
    add_run(p1, ".")
    set_paragraph_spacing(p1, space_after=8)

    # ── Section 2: How ──
    doc.add_heading("Section 2 — How", level=2)

    # Bullet 1: Return target exceeded
    b1 = doc.add_paragraph(style="List Bullet")
    add_bold_run(b1, "Return target exceeded. ")
    add_run(b1, "The blended 8.7% forecast return clears the Fund’s CPI + 2.5% nominal hurdle (5.2%) by 350 bps, with margin to absorb fee drag (~35 bps) and adverse rate scenarios (RBA, 2026).")
    set_paragraph_spacing(b1, space_after=4)

    # Bullet 2: Liquidity satisfied
    b2 = doc.add_paragraph(style="List Bullet")
    add_bold_run(b2, "Liquidity satisfied. ")
    add_run(b2, "STI (15%, daily redemption) covers the 10%-within-12-months requirement; STI plus MTG (50%) exceeds the 25%-within-3-years threshold, ensuring the Fund can meet drought-relief disbursements at short notice.")
    set_paragraph_spacing(b2, space_after=4)

    # Bullet 3: Risk appetite matched
    b3 = doc.add_paragraph(style="List Bullet")
    add_bold_run(b3, "Risk appetite matched. ")
    add_run(b3, "Total equity exposure reaches ~45%, consistent with the Board’s moderate-high tolerance. STI’s cash–bond anchor (0.7% volatility) floors downside; LTG’s 10-year horizon absorbs equity drawdowns.")
    set_paragraph_spacing(b3, space_after=8)

    # ── Section 3: Why ──
    doc.add_heading("Section 3 — Why", level=2)

    # Sub-heading: Macro regime
    h3a = doc.add_heading("Macro regime", level=3)
    set_paragraph_spacing(h3a, space_before=4, space_after=4)

    m1 = doc.add_paragraph(style="List Bullet")
    add_run(m1, "Australian headline CPI has moderated to 2.7%; the RBA signals gradual easing through 2026 (RBA, 2026). US inflation at 2.3%; the Fed maintains cautious normalisation (Federal Reserve, 2026). The IMF projects global growth at 3.3% as disinflation broadens (IMF, 2026).")
    set_paragraph_spacing(m1, space_after=4)

    m2 = doc.add_paragraph(style="List Bullet")
    add_run(m2, "Rate cuts support bond prices while equities benefit from improved earnings multiples. AUD depreciation on rate differentials favours unhedged global exposures.")
    set_paragraph_spacing(m2, space_after=6)

    # Sub-heading: Asset class outlook
    h3b = doc.add_heading("Asset class outlook", level=3)
    set_paragraph_spacing(h3b, space_before=4, space_after=4)

    a1 = doc.add_paragraph(style="List Bullet")
    add_run(a1, "Building-block estimates — combining current yields, consensus earnings growth, and policy-rate paths — produce forecasts respecting the risk-return principle: cash (3.8%) < bonds (3.9–4.8%) < property (8.4%) < equities (11.7–12.2%) < private equity (13.4%) (Bloomberg, 2026).")
    set_paragraph_spacing(a1, space_after=4)

    a2 = doc.add_paragraph(style="List Bullet")
    add_run(a2, "Easing supports fixed income: Australian FI forecasts 4.8% as yields compress; global FI (hedged) 3.9%. Global equities resist inflation drag through US corporate pricing power, technology-sector earnings resilience, and forward P/E expansion from rate cuts — forecasting 12.2% unhedged. Australian equities forecast 11.7% on domestic earnings recovery. Private equity leads at 13.4% but carries 20.3% volatility; infrastructure offers 8.6% with defensive, inflation-linked cash flows.")
    set_paragraph_spacing(a2, space_after=6)

    # Sub-heading: Unit trust logic
    h3c = doc.add_heading("Unit trust logic", level=3)
    set_paragraph_spacing(h3c, space_before=4, space_after=4)

    u1 = doc.add_paragraph(style="List Bullet")
    add_run(u1, "STI (50% cash / 50% short bond) blends to a forecast 4.1% at near-zero volatility — a pure liquidity vehicle. MTG’s 45% equity / 35% fixed income mix produces a forecast 7.9%. LTG’s 60% listed equities, 15% private equity, and 10% infrastructure drive the highest trust return at 10.6%. Each trust overweights asset classes forecast to outperform under disinflation.")
    set_paragraph_spacing(u1, space_after=6)

    # Sub-heading: Allocation rationale
    h3d = doc.add_heading("Allocation rationale", level=3)
    set_paragraph_spacing(h3d, space_before=4, space_after=4)

    ar1 = doc.add_paragraph(style="List Bullet")
    add_run(ar1, "The 15/35/50 split captures the disinflation tailwind: LTG (10.6%) maximises long-term real returns over a 10-year horizon; MTG (7.9%) balances growth and stability over 5 years; STI (4.1%) buffers near-term obligations. Blended: 0.15 × 4.1 + 0.35 × 7.9 + 0.50 × 10.6 = ")
    add_bold_run(ar1, "8.7%")
    add_run(ar1, " gross return at ~8.1% risk, clearing the 5.2% hurdle for the Fund’s moderate-high mandate.")
    set_paragraph_spacing(ar1, space_after=8)

    # ═══════════════════════════════════════════════════
    # PAGE BREAK
    # ═══════════════════════════════════════════════════
    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # PAGE 2
    # ═══════════════════════════════════════════════════

    pg2_title = doc.add_heading("Asset Class Review & Portfolio Metrics", level=2)
    set_paragraph_spacing(pg2_title, space_before=0, space_after=10)

    # ── Table 1: Asset Class Review ──
    t1_label = doc.add_paragraph()
    add_bold_run(t1_label, "Table 1: Asset Class Review", font_size=11)
    set_paragraph_spacing(t1_label, space_before=4, space_after=4)

    # Data for Table 1 — with forecast returns from final_content where available
    t1_headers = ["Asset Class", "Historical\nReturn", "Forecast\nReturn", "Diff.", "Historical\nRisk"]
    t1_data = [
        ["Cash",                                   "2.14%",  "3.8%",   "+1.66%", "0.43%"],
        ["Aus. Short Duration Bond",               "2.79%",  "—",     "—",    "1.24%"],
        ["Australian Fixed Income",                "2.11%",  "4.8%",   "+2.69%", "4.40%"],
        ["Global Fixed Income (Hedged)",           "2.08%",  "3.9%",   "+1.82%", "4.00%"],
        ["Global Credit (Hedged)",                 "2.68%",  "—",     "—",    "5.50%"],
        ["Australian Listed Equity",               "9.75%",  "11.7%",  "+1.95%", "13.47%"],
        ["Global Listed Equity (Unhedged)",        "10.14%", "12.2%",  "+2.06%", "14.32%"],
        ["Global Listed Equity (Hedged)",          "11.46%", "—",     "—",    "13.31%"],
        ["Australian Listed Property",             "7.36%",  "8.4%",   "+1.04%", "20.47%"],
        ["Global Infrastructure Unlisted (Unh.)",  "6.68%",  "8.6%",   "+1.92%", "10.35%"],
        ["Global Private Equity",                  "10.56%", "13.4%",  "+2.84%", "20.31%"],
    ]

    table1 = doc.add_table(rows=1 + len(t1_data), cols=5)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.autofit = True
    set_table_borders(table1)

    # Header row
    for i, h in enumerate(t1_headers):
        cell = table1.rows[0].cells[i]
        align = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        format_table_cell(cell, h, bold=True, align=align, font_size=9)
        set_cell_shading(cell, "1F3864")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for r_idx, row_data in enumerate(t1_data):
        for c_idx, val in enumerate(row_data):
            cell = table1.rows[r_idx + 1].cells[c_idx]
            align = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            format_table_cell(cell, val, bold=False, align=align, font_size=9)
            if r_idx % 2 == 0:
                set_cell_shading(cell, "E8EDF4")

    # Table 1 notes
    t1_notes = doc.add_paragraph()
    nr = t1_notes.add_run("Notes: Historical Return = CAGR, Jan 2016 – Feb 2026 (10-year). Historical Risk = annualised std. dev. of monthly returns. Forecast Return = 10-year forward building-block estimate. — = pending team input. Source: Refinitiv.")
    nr.font.size = Pt(8)
    nr.font.name = "Calibri"
    nr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    nr.italic = True
    set_paragraph_spacing(t1_notes, space_before=2, space_after=10)

    # ── Table 2: Unit Trust Performance ──
    t2_label = doc.add_paragraph()
    add_bold_run(t2_label, "Table 2: Unit Trust Performance", font_size=11)
    set_paragraph_spacing(t2_label, space_before=6, space_after=4)

    t2_headers = ["Metric", "STI", "MTG", "LTG"]
    t2_data = [
        ["Historical Return (net)", "2.28%", "5.91%", "8.39%"],
        ["Forecast Return (net)",   "4.1%",  "7.9%",  "10.6%"],
        ["Historical Risk",         "0.73%", "7.73%", "10.63%"],
    ]

    table2 = doc.add_table(rows=1 + len(t2_data), cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.autofit = True
    set_table_borders(table2)

    for i, h in enumerate(t2_headers):
        cell = table2.rows[0].cells[i]
        align = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        format_table_cell(cell, h, bold=True, align=align, font_size=9)
        set_cell_shading(cell, "1F3864")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row_data in enumerate(t2_data):
        for c_idx, val in enumerate(row_data):
            cell = table2.rows[r_idx + 1].cells[c_idx]
            align = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            format_table_cell(cell, val, bold=False, align=align, font_size=9)
            if r_idx % 2 == 0:
                set_cell_shading(cell, "E8EDF4")

    # Table 2 notes
    t2_notes = doc.add_paragraph()
    nr2 = t2_notes.add_run("Notes: Historical Return = weighted-average CAGR net of costs. Forecast Return = weighted-average of asset-class forecasts. Risk = annualised std. dev. of trust monthly returns (captures correlations). Trust compositions per Abercrombie FM IM Appendix 2. Source: Refinitiv, team analysis.")
    nr2.font.size = Pt(8)
    nr2.font.name = "Calibri"
    nr2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    nr2.italic = True
    set_paragraph_spacing(t2_notes, space_before=2, space_after=10)

    # ── Table 3: NSWDF Portfolio ──
    t3_label = doc.add_paragraph()
    add_bold_run(t3_label, "Table 3: NSWDF Portfolio Recommendation", font_size=11)
    set_paragraph_spacing(t3_label, space_before=6, space_after=4)

    t3_headers = ["Metric", "Value"]
    t3_data = [
        ["Recommended Mix", "STI: 15%    MTG: 35%    LTG: 50%"],
        ["Forecast Return (gross)", "8.7% p.a."],
        ["Estimated Portfolio Risk", "~8.1%"],
        ["CPI + 2.5% Hurdle", "5.2% (exceeded by 350 bps)"],
        ["Liquidity Coverage", "STI 15% (>10% at 12 mo.); STI+MTG 50% (>25% at 3 yr)"],
    ]

    table3 = doc.add_table(rows=1 + len(t3_data), cols=2)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    table3.autofit = True
    set_table_borders(table3)

    for i, h in enumerate(t3_headers):
        cell = table3.rows[0].cells[i]
        format_table_cell(cell, h, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=9)
        set_cell_shading(cell, "1F3864")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row_data in enumerate(t3_data):
        for c_idx, val in enumerate(row_data):
            cell = table3.rows[r_idx + 1].cells[c_idx]
            bold = (c_idx == 1 and r_idx in [0, 1])  # Bold the key values
            format_table_cell(cell, val, bold=bold, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=9)
            if r_idx % 2 == 0:
                set_cell_shading(cell, "E8EDF4")

    # Table 3 notes
    t3_notes = doc.add_paragraph()
    nr3 = t3_notes.add_run("Notes: Forecast Return = 0.15×4.1 + 0.35×7.9 + 0.50×10.6 = 8.7%. Risk estimated from historical trust covariance at recommended weights. Constraints: CPI+2.5% return hurdle, 10%/12-mo and 25%/3-yr liquidity, moderate-high risk tolerance. Fund size AUD 3 bn. Source: Team analysis; Board Policy and Investment Directive.")
    nr3.font.size = Pt(8)
    nr3.font.name = "Calibri"
    nr3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    nr3.italic = True
    set_paragraph_spacing(t3_notes, space_before=2, space_after=16)

    # ── AI Acknowledgement Statement ──
    divider = doc.add_paragraph()
    dr = divider.add_run("─" * 60)
    dr.font.size = Pt(8)
    dr.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
    set_paragraph_spacing(divider, space_before=8, space_after=4)

    ai_heading = doc.add_paragraph()
    add_bold_run(ai_heading, "AI Acknowledgement", font_size=10)
    set_paragraph_spacing(ai_heading, space_before=0, space_after=4)

    ai_text = doc.add_paragraph()
    ai_run = ai_text.add_run(
        "This briefing paper was prepared with the assistance of generative AI tools "
        "(Claude, Anthropic). AI was used to: (1) extract and validate data from the "
        "source Excel workbook, (2) compute historical returns, risk metrics, and unit "
        "trust performance figures, (3) structure and draft the narrative content, and "
        "(4) format the final document. All figures, recommendations, and analytical "
        "judgements were reviewed and approved by the student team. The underlying data "
        "originates from Refinitiv indices as provided in the FINC3600 project brief."
    )
    ai_run.font.size = Pt(9)
    ai_run.font.name = "Calibri"
    ai_run.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
    ai_run.italic = True
    set_paragraph_spacing(ai_text, space_before=0, space_after=0)

    # ── Save ──
    out_path = "/Users/sty/OneDrive - The University of Sydney (Students)/QBUS3600/qbus3600_abs-zombies/agent-system/outputs/CFO_Briefing_NSW_Drought_Fund.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    create_cfo_brief()
