#!/usr/bin/env python3
"""Build CFO Briefing .docx — Page 1 narrative + Page 2 data appendix."""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Output path ──────────────────────────────────────────────────────────────
OUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(OUT_DIR, "CFO_Briefing_NSW_Drought_Fund.docx")

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    """Set cell background colour."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, font_size=Pt(10), alignment=None, font_color=None):
    """Write formatted text into a table cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.size = font_size
    run.font.name = "Calibri"
    run.bold = bold
    if font_color:
        run.font.color.rgb = font_color
    # Tight cell spacing
    fmt = p.paragraph_format
    fmt.space_before = Pt(1)
    fmt.space_after = Pt(1)
    fmt.line_spacing = Pt(12)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table with distinct header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, bold=True, font_size=Pt(9),
                      alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      font_color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(hdr_cells[i], "1F3864")  # dark navy

    # Data rows
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[c_idx], val, font_size=Pt(9), alignment=align)
            # Alternating row shading
            if r_idx % 2 == 1:
                set_cell_shading(cells[c_idx], "D6E4F0")

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w

    return table


# ── Document setup ───────────────────────────────────────────────────────────
doc = Document()

# -- Default font & spacing --------------------------------------------------
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(4)

# Heading styles
for level, size, color in [
    ("Heading 1", 16, "1F3864"),
    ("Heading 2", 13, "2E75B6"),
    ("Heading 3", 12, "2E75B6"),
]:
    hs = doc.styles[level]
    hs.font.name = "Calibri"
    hs.font.size = Pt(size)
    hs.font.color.rgb = RGBColor.from_string(color)
    hs.font.bold = True
    hs.paragraph_format.space_before = Pt(10)
    hs.paragraph_format.space_after = Pt(4)
    hs.paragraph_format.line_spacing = 1.15

# -- Page layout: portrait, 2.54 cm margins ----------------------------------
section = doc.sections[0]
section.orientation = WD_ORIENT.PORTRAIT
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

# -- Footer: page numbers ----------------------------------------------------
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run._r.append(fld_char_begin)
run2 = fp.add_run()
instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
run2._r.append(instr)
run3 = fp.add_run()
fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run3._r.append(fld_char_end)

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — NARRATIVE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("CFO Briefing — NSW Drought Fund: Initial Portfolio Recommendation", level=1)

# ── Section 1: Recommendations ──────────────────────────────────────────────
doc.add_heading("Section 1: Recommendations", level=2)

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.15
run = p.add_run(
    "Under a disinflationary base case (Australian CPI 2.7%, RBA easing 75–100 bps "
    "over 2026–27), we recommend "
)
run.font.size = Pt(12)
run.font.name = "Calibri"

run_b = p.add_run("18% STI / 22% MTG / 60% LTG")
run_b.bold = True
run_b.font.size = Pt(12)
run_b.font.name = "Calibri"

run2 = p.add_run(
    " — tilting toward long-term growth to maximise real returns over the Fund's "
    "10-year horizon. The optimised portfolio delivers a forecast "
)
run2.font.size = Pt(12)
run2.font.name = "Calibri"

run_b2 = p.add_run("4.86% p.a. net return at 7.71% risk")
run_b2.bold = True
run_b2.font.size = Pt(12)
run_b2.font.name = "Calibri"

run3 = p.add_run(" (Table 3).")
run3.font.size = Pt(12)
run3.font.name = "Calibri"

# ── Section 2: Key Implications ─────────────────────────────────────────────
doc.add_heading("Section 2: Key Implications", level=2)

implications = [
    (
        "Return vs hurdle: ",
        "The 4.86% net return falls 34 bps short of the CPI + 2.5% hurdle (5.2%) under "
        "CMA v2 building-block assumptions. This reflects deliberately conservative equity "
        "forecasts (5.4–8.2% vs 9.8–11.5% historical; Table 1). Any equity outperformance "
        "relative to CMA v2 clears the hurdle; the 60% LTG tilt maximises that probability "
        "over the Fund's intergenerational investment horizon (RBA, 2026)."
    ),
    (
        "Liquidity: ",
        "STI (18%) exceeds the 10%-within-12-months requirement; STI plus MTG (40%) exceeds "
        "the 25%-within-3-years threshold. Both mandates are satisfied with margin."
    ),
    (
        "Risk appetite: ",
        "Portfolio risk of 7.71% sits within the Board's moderate-high tolerance. STI anchors "
        "near-term volatility at 0.90%; LTG's 10-year horizon absorbs short-term drawdowns (Table 2)."
    ),
]

for bold_prefix, body in implications:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.15
    run_b = p.add_run(bold_prefix)
    run_b.bold = True
    run_b.font.size = Pt(12)
    run_b.font.name = "Calibri"
    run_t = p.add_run(body)
    run_t.font.size = Pt(12)
    run_t.font.name = "Calibri"

# ── Section 3: Key Findings ────────────────────────────────────────────────
doc.add_heading("Section 3: Key Findings", level=2)

# A. Macro regime
doc.add_heading("A. Macro Regime", level=3)

macro_bullets = [
    (
        "Australia: ",
        "CPI at 2.7% sits within the RBA's 2–3% target band. The anticipated easing cycle "
        "compresses bond yields — generating fixed-income capital gains — and reduces corporate "
        "borrowing costs, supporting equity valuations (RBA, 2026)."
    ),
    (
        "Global: ",
        "US CPI at 2.3%; IMF projects global growth at 3.3% (IMF, 2026). AUD depreciation of "
        "~5% YTD boosts unhedged global equity and infrastructure returns for AUD-based investors."
    ),
]

for bold_prefix, body in macro_bullets:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.15
    run_b = p.add_run(bold_prefix)
    run_b.bold = True
    run_b.font.size = Pt(12)
    run_b.font.name = "Calibri"
    run_t = p.add_run(body)
    run_t.font.size = Pt(12)
    run_t.font.name = "Calibri"

# B. Asset class outlook
doc.add_heading("B. Asset Class Outlook (NSWDF CMA v2, Building-Block)", level=3)

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.15
run = p.add_run(
    "Fixed-income forecasts exceed 10-year history by 1.2–2.3 pp because 2026 starting yields "
    "sit above the decade average. Equity forecasts fall 2.4–5.3 pp below history due to elevated "
    "starting valuations compressing forward returns (Table 1). This asymmetry drives the "
    "conservative net-return estimate and motivates the LTG overweight to capture equity upside "
    "should valuations normalise."
)
run.font.size = Pt(12)
run.font.name = "Calibri"

# C. Unit trust construction
doc.add_heading("C. Unit Trust Construction", level=3)

trust_bullets = [
    (
        "STI ",
        "(50% cash / 50% short bonds → 3.57% net, 0.90% vol) is a pure liquidity vehicle."
    ),
    (
        "MTG ",
        "(45% equity / 35% FI → 4.67% net) balances growth and stability."
    ),
    (
        "LTG ",
        "(60% equities / 15% PE / 10% infrastructure → 5.31% net) overweights assets forecast "
        "to outperform under disinflation and rate easing (Table 2)."
    ),
]

for bold_prefix, body in trust_bullets:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.15
    run_b = p.add_run(bold_prefix)
    run_b.bold = True
    run_b.font.size = Pt(12)
    run_b.font.name = "Calibri"
    run_t = p.add_run(body)
    run_t.font.size = Pt(12)
    run_t.font.name = "Calibri"

# D. Allocation rationale
doc.add_heading("D. Allocation Rationale", level=3)

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.15
run = p.add_run("0.18 × 3.57 + 0.22 × 4.67 + 0.60 × 5.31 = ")
run.font.size = Pt(12)
run.font.name = "Calibri"

run_b = p.add_run("4.86% net")
run_b.bold = True
run_b.font.size = Pt(12)
run_b.font.name = "Calibri"

run2 = p.add_run(
    " at 7.71% risk (Table 3). LTG maximises real returns; MTG provides growth-stability "
    "balance; STI buffers near-term obligations."
)
run2.font.size = Pt(12)
run2.font.name = "Calibri"

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE BREAK → PAGE 2 — DATA APPENDIX
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

doc.add_heading("Data Appendix", level=1)

# ── Table 1: Asset Class Review ─────────────────────────────────────────────
doc.add_heading("Table 1: Asset Class Review", level=2)

t1_headers = ["Asset Class", "Historical Returns", "Forecast Returns", "Difference", "Historical Risk"]
t1_rows = [
    ["Cash",                                        "2.14%", "3.35%", "+1.21%", "0.43%"],
    ["Australian Short Duration Bond",              "2.79%", "4.15%", "+1.36%", "1.24%"],
    ["Australian Fixed Income",                     "2.11%", "4.35%", "+2.24%", "4.40%"],
    ["Global Fixed Income (Hedged)",                "2.08%", "4.15%", "+2.07%", "4.00%"],
    ["Global Credit (Hedged)",                      "2.68%", "5.00%", "+2.32%", "5.50%"],
    ["Australian Listed Equity",                    "9.75%", "5.40%", "−4.35%", "13.47%"],
    ["Global Listed Equity (Unhedged)",             "10.14%", "6.35%", "−3.79%", "14.32%"],
    ["Global Listed Equity (Hedged)",               "11.46%", "6.20%", "−5.26%", "13.31%"],
    ["Australian Listed Property",                  "7.36%", "6.80%", "−0.56%", "20.47%"],
    ["Global Infrastructure Unlisted (Unhedged)",   "6.68%", "6.90%", "+0.22%", "10.35%"],
    ["Global Private Equity",                       "10.56%", "8.20%", "−2.36%", "20.31%"],
]

add_styled_table(doc, t1_headers, t1_rows, col_widths=[Cm(5.5), Cm(2.8), Cm(2.8), Cm(2.2), Cm(2.4)])

# Table 1 notes
p_note = doc.add_paragraph()
p_note.paragraph_format.space_before = Pt(4)
p_note.paragraph_format.space_after = Pt(2)
p_note.paragraph_format.line_spacing = 1.0
run_n = p_note.add_run(
    "Notes. Historical Returns = annualised geometric mean (CAGR) of monthly index total returns, "
    "Jan 2016 – Feb 2026. Forecast Returns = 10-year gross compound return from building-block "
    "decomposition. Difference = Forecast minus Historical. Historical Risk = annualised standard "
    "deviation of monthly returns. Source: Historical — Refinitiv via Bloomberg proxy indices; "
    "Forecast — NSWDF 10Y CMA v2 (May 2026)."
)
run_n.font.size = Pt(8)
run_n.font.name = "Calibri"
run_n.italic = True
run_n.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

# ── Table 2: Unit Trust Performance ─────────────────────────────────────────
doc.add_heading("Table 2: Unit Trust Performance", level=2)

t2_headers = ["Metric", "STI", "MTG", "LTG"]
t2_rows = [
    ["Expected Return", "3.57%", "4.67%", "5.31%"],
    ["Risk",            "0.90%", "6.79%", "10.54%"],
]

add_styled_table(doc, t2_headers, t2_rows, col_widths=[Cm(4.0), Cm(3.5), Cm(3.5), Cm(3.5)])

p_note2 = doc.add_paragraph()
p_note2.paragraph_format.space_before = Pt(4)
p_note2.paragraph_format.space_after = Pt(2)
p_note2.paragraph_format.line_spacing = 1.0
run_n2 = p_note2.add_run(
    "Notes. Expected Return = weighted-average of 10-year forecast returns (Table 1) across each "
    "trust's underlying asset classes, net of asset-class proxy costs and trust-level ongoing fees. "
    "Risk = annualised standard deviation of the trust's portfolio return. Trust compositions per "
    "IM Table 4A. Source: Computed from NSWDF 10Y CMA v2 forecasts, IM Table 4A compositions, "
    "and IM Tables 3A/4A cost schedule."
)
run_n2.font.size = Pt(8)
run_n2.font.name = "Calibri"
run_n2.italic = True
run_n2.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

# ── Table 3: NSWDF Portfolio ────────────────────────────────────────────────
doc.add_heading("Table 3: NSWDF Portfolio", level=2)

t3_headers = ["Metric", "Value"]
t3_rows = [
    ["Recommended Mix of Unit Trusts", "STI 18% / MTG 22% / LTG 60%"],
    ["Forecast Return",                "4.86% p.a."],
    ["Forecast Risk",                  "7.71% p.a."],
]

add_styled_table(doc, t3_headers, t3_rows, col_widths=[Cm(6.0), Cm(8.0)])

p_note3 = doc.add_paragraph()
p_note3.paragraph_format.space_before = Pt(4)
p_note3.paragraph_format.space_after = Pt(2)
p_note3.paragraph_format.line_spacing = 1.0
run_n3 = p_note3.add_run(
    "Notes. The recommended mix maximises forecast net return subject to: (i) portfolio volatility "
    "capped at 10% p.a., (ii) STI ≥ 15% to meet the 10%-within-12-months liquidity call plus a 5% "
    "buffer, (iii) STI + MTG ≥ 40% to meet the 25%-within-3-years call plus buffer, (iv) LTG ≤ 60% "
    "concentration cap, and (v) parametric 95% VaR ≤ 10%. Fund size: AUD 3 billion; Board Policy "
    "target: CPI + 2.5% = 5.1% p.a. over a rolling 10-year period. Source: Constrained optimisation "
    "(SciPy SLSQP) using NSWDF 10Y CMA v2 inputs; Board Policy and Investment Directive."
)
run_n3.font.size = Pt(8)
run_n3.font.name = "Calibri"
run_n3.italic = True
run_n3.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

# ── AI Acknowledgement ──────────────────────────────────────────────────────
doc.add_paragraph()  # spacer
p_ai = doc.add_paragraph()
p_ai.paragraph_format.space_before = Pt(12)
p_ai.paragraph_format.line_spacing = 1.15

run_ai_title = p_ai.add_run("AI Acknowledgement Statement")
run_ai_title.bold = True
run_ai_title.font.size = Pt(11)
run_ai_title.font.name = "Calibri"

p_ai2 = doc.add_paragraph()
p_ai2.paragraph_format.line_spacing = 1.15
run_ai2 = p_ai2.add_run(
    "This briefing document was prepared with the assistance of generative AI tools "
    "(Claude, Anthropic). AI was used to structure content, draft narrative text, and "
    "format tables from source data. All quantitative inputs (capital market assumptions, "
    "historical returns, optimisation outputs) were produced independently by the project "
    "team using documented methodologies. The final document was reviewed, verified, and "
    "approved by the team prior to submission."
)
run_ai2.font.size = Pt(10)
run_ai2.font.name = "Calibri"
run_ai2.italic = True
run_ai2.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f"✅ Saved → {OUT_PATH}")
